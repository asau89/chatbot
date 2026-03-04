"""Multi-format document loader — PDF, DOCX, XLSX, TXT.

Extends the existing layout-aware PDF loader with support for additional
file formats commonly found in enterprise knowledge bases.  Each loader
returns a list of LangChain ``Document`` objects — one per logical page
or sheet — so all downstream processing (splitting, embedding, storage)
works unchanged.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from langchain_core.documents import Document

# PDF loader is the existing one
from rag.document_loader import load_pdf

# ---------------------------------------------------------------------------
# DOCX Loader
# ---------------------------------------------------------------------------

def load_docx(
    docx_path: str | Path,
    *,
    label: str = "",
) -> list[Document]:
    """Load a .docx file and return one Document per page-break section.

    Falls back to one document for the whole file if no page breaks are found.
    """
    from docx import Document as DocxDocument
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    doc = DocxDocument(str(docx_path))
    pages: list[str] = []
    current_parts: list[str] = []

    for para in doc.paragraphs:
        # Detect page breaks
        has_page_break = False
        for run in para.runs:
            if run._element.xml and "w:br" in run._element.xml and 'w:type="page"' in run._element.xml:
                has_page_break = True
                break

        if has_page_break and current_parts:
            pages.append("\n".join(current_parts))
            current_parts = []

        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                prefix = "##" if level in ("1", "2") else "###"
                current_parts.append(f"\n{prefix} {text}\n")
            else:
                current_parts.append(text)

    if current_parts:
        pages.append("\n".join(current_parts))

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            header_sep = "| " + " | ".join("---" for _ in table.rows[0].cells) + " |"
            table_md = rows[0] + "\n" + header_sep + "\n" + "\n".join(rows[1:])
            if pages:
                pages[-1] += f"\n\n{table_md}"
            else:
                pages.append(table_md)

    if not pages:
        return []

    final_label = label.strip() if label.strip() else docx_path.stem
    return [
        Document(
            page_content=text,
            metadata={
                "source": docx_path.name,
                "page": i,
                "extractor": "docx",
                "label": final_label,
            },
        )
        for i, text in enumerate(pages, start=1)
        if text.strip()
    ]


# ---------------------------------------------------------------------------
# XLSX Loader
# ---------------------------------------------------------------------------

def load_xlsx(
    xlsx_path: str | Path,
    *,
    label: str = "",
) -> list[Document]:
    """Load a .xlsx file and return one Document per worksheet.

    Each sheet is converted to a Markdown table.
    """
    from openpyxl import load_workbook

    xlsx_path = Path(xlsx_path)
    if not xlsx_path.exists():
        raise FileNotFoundError(f"XLSX not found: {xlsx_path}")

    wb = load_workbook(str(xlsx_path), read_only=True, data_only=True)
    documents: list[Document] = []
    final_label = label.strip() if label.strip() else xlsx_path.stem

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # Build Markdown table
        md_lines: list[str] = []
        header = rows[0]
        header_strs = [str(c or "").strip() for c in header]
        md_lines.append("| " + " | ".join(header_strs) + " |")
        md_lines.append("| " + " | ".join("---" for _ in header) + " |")

        for row in rows[1:]:
            cells = [str(c or "").strip() for c in row]
            # Skip empty rows
            if not any(cells):
                continue
            md_lines.append("| " + " | ".join(cells) + " |")

        text = f"## {sheet_name}\n\n" + "\n".join(md_lines)
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": xlsx_path.name,
                        "page": sheet_name,
                        "extractor": "xlsx",
                        "label": final_label,
                    },
                )
            )

    wb.close()
    return documents


# ---------------------------------------------------------------------------
# TXT Loader
# ---------------------------------------------------------------------------

def load_txt(
    txt_path: str | Path,
    *,
    label: str = "",
) -> list[Document]:
    """Load a plain-text file as a single Document."""
    txt_path = Path(txt_path)
    if not txt_path.exists():
        raise FileNotFoundError(f"TXT not found: {txt_path}")

    text = txt_path.read_text(encoding="utf-8", errors="replace")
    if not text.strip():
        return []

    final_label = label.strip() if label.strip() else txt_path.stem
    return [
        Document(
            page_content=text,
            metadata={
                "source": txt_path.name,
                "page": 1,
                "extractor": "txt",
                "label": final_label,
            },
        )
    ]


# ---------------------------------------------------------------------------
# Unified dispatcher
# ---------------------------------------------------------------------------

_LOADERS: dict[str, Any] = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
    ".xls": load_xlsx,   # openpyxl can read .xls via xlrd fallback
    ".txt": load_txt,
    ".md": load_txt,
    ".csv": load_txt,    # CSV treated as plain text — splits well downstream
}

SUPPORTED_EXTENSIONS = set(_LOADERS.keys())
SUPPORTED_MIMETYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "text/plain",
    "text/markdown",
    "text/csv",
}


def load_document(
    file_path: str | Path,
    *,
    label: str = "",
) -> list[Document]:
    """Auto-detect file type and load using the appropriate loader.

    Raises ``ValueError`` for unsupported file types.
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return loader(file_path, label=label)
